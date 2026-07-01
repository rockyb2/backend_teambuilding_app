from smolagents import Tool
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from copy import copy
from pathlib import Path

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, Reference
from openpyxl.drawing.image import Image
import os
import json

from email import encoders
from email.header import Header
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formataddr
import os
import smtplib
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class BuildExcelPro(Tool):
    name = "BuildExcelPro"
    description = (
        "Crée un fichier Excel professionnel avec tableau, formules, graphiques et image."
    )

    inputs = {
        "name": {
            "type": "string",
            "description": "Nom du fichier (sans .xlsx)",
            
        },
        "headers": {
            "type": "array",
            "description": "Liste des titres de colonnes",
            
        },
        "rows": {
            "type": "array",
            "description": "Toutes les données ligne par ligne",
            
        },
        # "chart": {
        #     "type": "boolean",
        #     "description": "Ajouter un graphique ?",
            
        # },
        # "image_path": {
        #     "type": "string",
        #     "description": "Chemin d'une image à insérer",
            
        # }
    }

    output_type = "string"

    cotation_headers = {
        "categorie",
        "poste",
        "description",
        "quantite",
        "prix unitaire estimatif",
        "total estimatif",
        "notes",
    }
    template_path = (
        Path(__file__).resolve().parents[1]
        / "assets"
        / "proforma"
        / "cotation_template.xlsx"
    )

    def _is_cotation(self, headers):
        normalized_headers = {str(header or "").strip().lower() for header in headers or []}
        return self.cotation_headers.issubset(normalized_headers)

    def _save_path(self, name: str) -> Path:
        file_path = Path(f"{name}.xlsx")
        file_path.parent.mkdir(parents=True, exist_ok=True)
        return file_path

    def _row_dict(self, headers, row):
        return {
            str(header or "").strip().lower(): row[index] if index < len(row) else ""
            for index, header in enumerate(headers)
        }

    def _as_number(self, value, default=0):
        if isinstance(value, bool):
            return default
        if isinstance(value, (int, float)):
            return value
        text = str(value or "").replace(" ", "").replace(",", ".")
        try:
            return float(text)
        except ValueError:
            return default

    def _clean_text(self, value, fallback="A completer"):
        text = str(value or "").strip()
        return text if text else fallback

    def _metadata_and_items(self, headers, rows):
        metadata = {}
        items = []

        for row in rows or []:
            row_data = self._row_dict(headers, row)
            category = self._clean_text(row_data.get("categorie"), "")
            post = self._clean_text(row_data.get("poste"), "")

            if category.lower() == "dossier":
                metadata[post.lower()] = self._clean_text(row_data.get("description"), "")
                continue
            if category.lower() == "total":
                continue

            items.append(
                {
                    "category": category or "Cotation",
                    "post": post or "Poste a completer",
                    "description": self._clean_text(row_data.get("description"), ""),
                    "quantity": self._as_number(row_data.get("quantite"), 1) or 1,
                    "unit_price": self._as_number(row_data.get("prix unitaire estimatif"), 0),
                    "total": self._as_number(row_data.get("total estimatif"), 0),
                    "notes": self._clean_text(row_data.get("notes"), ""),
                }
            )

        return metadata, items

    def _copy_cell_style(self, source, target):
        target.font = copy(source.font)
        target.fill = copy(source.fill)
        target.border = copy(source.border)
        target.alignment = copy(source.alignment)
        target.number_format = source.number_format
        target.protection = copy(source.protection)

    def _copy_row_style(self, ws, source_row, target_row, min_col=2, max_col=7):
        ws.row_dimensions[target_row].height = ws.row_dimensions[source_row].height
        for col_index in range(min_col, max_col + 1):
            self._copy_cell_style(ws.cell(source_row, col_index), ws.cell(target_row, col_index))

    def _fill_merged_row(self, ws, row_index, value, source_row=16):
        ws.merge_cells(start_row=row_index, start_column=2, end_row=row_index, end_column=7)
        self._copy_row_style(ws, source_row, row_index)
        cell = ws.cell(row_index, 2)
        cell.value = value
        cell.alignment = Alignment(horizontal="center", vertical="center")

    def _write_item_row(self, ws, row_index, item, source_row=17):
        self._copy_row_style(ws, source_row, row_index)
        label = item["post"]
        if item["description"] and item["description"].lower() != label.lower():
            label = f"{label} - {item['description']}"

        ws.cell(row_index, 2).value = label
        ws.cell(row_index, 3).value = item["quantity"]
        ws.cell(row_index, 4).value = item["unit_price"]
        ws.cell(row_index, 5).value = 1
        ws.cell(row_index, 6).value = f"=C{row_index}*D{row_index}*E{row_index}"
        ws.cell(row_index, 7).value = item["notes"]

        ws.cell(row_index, 2).alignment = Alignment(wrap_text=True, vertical="center")
        ws.cell(row_index, 3).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row_index, 5).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row_index, 6).number_format = '#,##0 "CFA"'

    def _write_subtotal_row(self, ws, row_index, category, first_row, last_row, source_row=19):
        self._copy_row_style(ws, source_row, row_index)
        ws.cell(row_index, 3).value = f"SOUS-TOTAL {category.upper()}"
        ws.cell(row_index, 6).value = f"=SUM(F{first_row}:F{last_row})"
        ws.cell(row_index, 6).number_format = '#,##0 "CFA"'
        return ws.cell(row_index, 6).coordinate

    def _build_cotation_from_template(self, name, headers, rows):
        if not self.template_path.is_file():
            return None

        metadata, items = self._metadata_and_items(headers, rows)
        if not items:
            return None

        wb = load_workbook(self.template_path)
        ws = wb.worksheets[0]
        ws.title = "COTATION"
        for sheet in list(wb.worksheets)[1:]:
            wb.remove(sheet)

        ws["C9"] = metadata.get("date/periode") or metadata.get("date") or "=TODAY()"
        ws["F9"] = metadata.get("lieu") or metadata.get("lieu souhaite") or "A completer"
        ws["C10"] = metadata.get("client") or "Client a completer"
        ws["F10"] = metadata.get("duree") or 1
        ws["C11"] = (
            self._as_number(metadata.get("participants"), metadata.get("nombre_personnes") or 0)
            or "A completer"
        )
        ws["F11"] = metadata.get("nuits") or ""
        ws["G15"] = "Notes"
        ws.column_dimensions["G"].width = 34
        self._copy_cell_style(ws["F15"], ws["G15"])

        template_max_row = ws.max_row
        for merged_range in list(ws.merged_cells.ranges):
            if merged_range.min_row >= 16:
                ws.unmerge_cells(str(merged_range))
        for row in range(16, template_max_row + 1):
            for col in range(2, 8):
                ws.cell(row, col).value = None

        grouped_items = []
        for item in items:
            if not grouped_items or grouped_items[-1][0] != item["category"]:
                grouped_items.append((item["category"], []))
            grouped_items[-1][1].append(item)

        row_index = 16
        subtotal_cells = []
        for category, category_items in grouped_items:
            self._fill_merged_row(ws, row_index, category.upper(), source_row=15)
            row_index += 1
            first_item_row = row_index

            for item in category_items:
                self._write_item_row(ws, row_index, item)
                row_index += 1

            subtotal_cells.append(
                self._write_subtotal_row(
                    ws,
                    row_index,
                    category,
                    first_item_row,
                    row_index - 1,
                )
            )
            row_index += 2

        self._write_subtotal_row(ws, row_index, "MISE EN OEUVRE", row_index, row_index)
        ws.cell(row_index, 6).value = f"=SUM({','.join(subtotal_cells)})" if subtotal_cells else 0
        implementation_total_cell = ws.cell(row_index, 6).coordinate
        row_index += 2

        self._copy_row_style(ws, 19, row_index)
        ws.cell(row_index, 4).value = "TOTAL ESTIMATIF"
        ws.cell(row_index, 6).value = f"={implementation_total_cell}"
        ws.cell(row_index, 6).number_format = '#,##0 "CFA"'
        row_index += 2

        self._fill_merged_row(ws, row_index, "NOTES INTERNES", source_row=15)
        row_index += 1
        self._copy_row_style(ws, 17, row_index)
        ws.merge_cells(start_row=row_index, start_column=2, end_row=row_index, end_column=7)
        notes = [
            "Cotation brouillon interne a verifier avant envoi au client.",
            "Les montants restent indicatifs et doivent etre valides par l'equipe commerciale.",
        ]
        budget = metadata.get("budget") or metadata.get("budget estime")
        if budget:
            notes.append(f"Budget indique par le client: {budget}")
        ws.cell(row_index, 2).value = "\n".join(notes)
        ws.cell(row_index, 2).alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[row_index].height = 58

        ws.print_area = f"B1:G{row_index}"
        ws.sheet_view.showGridLines = False
        if hasattr(wb, "calculation"):
            wb.calculation.fullCalcOnLoad = True

        file_path = self._save_path(name)
        wb.save(file_path)
        return str(file_path)

    def _build_simple_workbook(self, name, headers, rows):
        wb = Workbook()
        ws = wb.active
        ws.title = "Feuille1"

        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="4472C4")
            cell.alignment = Alignment(horizontal="center")

        for row in rows:
            ws.append(row)

        last_row = ws.max_row
        last_col = ws.max_column
        table = Table(
            displayName="Table1",
            ref=f"A1:{get_column_letter(last_col)}{last_row}",
        )
        style = TableStyleInfo(
            name="TableStyleMedium9",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False
        )
        table.tableStyleInfo = style
        ws.add_table(table)

        for col in ws.columns:
            max_length = max(len(str(c.value)) for c in col)
            ws.column_dimensions[col[0].column_letter].width = max_length + 2

        file_path = self._save_path(name)
        wb.save(file_path)
        return str(file_path)

    def forward(self, name, headers, rows, ):
        try:
            file_name = None
            if self._is_cotation(headers):
                file_name = self._build_cotation_from_template(name, headers, rows)
            if not file_name:
                file_name = self._build_simple_workbook(name, headers, rows)

            return f"Excel '{file_name}' genere avec succes !||{file_name}"

        except Exception as e:
            return f"Erreur Excel Pro : {e}"


class BuildPDF(Tool):
    name = "BuildPDF"
    description = "Génère un PDF professionnel avec titre, contenu, marges et styles optimisés."

    inputs = {
        "name": {"type": "string", "description": "Nom du fichier PDF sans extension"},
        "title": {"type": "string", "description": "Titre du PDF"},
        "content": {"type": "string", "description": "Texte du PDF"},
    }
    output_type = "string"

    def forward(self, name: str, title: str, content: str) -> str:
        try:
            file_name = f"{name}.pdf"
            styles = getSampleStyleSheet()
            style_title = styles["Title"]
            style_body = styles["BodyText"]

            doc = SimpleDocTemplate(
                file_name,
                pagesize=A4,
                leftMargin=2*cm,
                rightMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )

            story = [
                Paragraph(title, style_title),
                Spacer(1, 12),
                Paragraph(content.replace("\n", "<br/>"), style_body)
            ]

            doc.build(story)

            return f"PDF '{file_name}' généré avec succès||{file_name}"

        except Exception as e:
            return f"Erreur PDF : {str(e)}"




class BuildWord(Tool):
    name = "BuildWord"
    description = (
        "Crée un document Word professionnel (.docx) avec support avancé des tableaux. "
        "Peut générer des lettres, CV, rapports avec tableaux simples ou complexes. "
        "Supporte la fusion de cellules, styles personnalisés, et mise en page avancée."
    )

    inputs = {
        "title": {"type": "string", "description": "Titre principal du document"},
        "recipient": {"type": "string", "description": "Destinataire (optionnel, multiligne)", "nullable": True},
        "sender": {"type": "string", "description": "Expéditeur (optionnel, multiligne)", "nullable": True},
        "date": {"type": "string", "description": "Date (optionnel)", "nullable": True},
        "subject": {"type": "string", "description": "Objet (optionnel)", "nullable": True},
        "body": {"type": "string", "description": "Corps du document. Supporte :\\n- Paragraphes (séparés par \\n\\n)\\n- Listes (lignes commençant par '- ')\\n- Tableaux (JSON: [TABLE]json_data[/TABLE])", "nullable": True},
        "tables": {"type": "string", "description": "JSON array de tableaux complexes (optionnel). Format: [{headers:[...], rows:[[...]], merge_cells:[...], styles:{...}}]", "nullable": True},
        "filename": {"type": "string", "description": "Nom du fichier sans extension," ,  "nullable": True},
        "style_config": {"type": "string", "description": "JSON de config style (optionnel): {font:'Arial', font_size:11, margins:2.5, line_spacing:1.15}", "nullable": True}
    }

    output_type = "string"

    def _set_cell_border(self, cell, **kwargs):
        """Ajoute des bordures personnalisées à une cellule"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                edge_data = kwargs[edge]
                edge_el = OxmlElement(f'w:{edge}')
                edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
                edge_el.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                edge_el.set(qn('w:space'), str(edge_data.get('space', 0)))
                edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
                tcBorders.append(edge_el)
        
        tcPr.append(tcBorders)

    def _apply_cell_style(self, cell, style_config):
        """Applique un style à une cellule"""
        if 'background' in style_config:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), style_config['background'])
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        if 'bold' in style_config or 'font_size' in style_config or 'color' in style_config:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if style_config.get('bold'):
                        run.font.bold = True
                    if 'font_size' in style_config:
                        run.font.size = Pt(style_config['font_size'])
                    if 'color' in style_config:
                        color = style_config['color']
                        run.font.color.rgb = RGBColor(
                            int(color[0:2], 16),
                            int(color[2:4], 16),
                            int(color[4:6], 16)
                        )
        
        if 'alignment' in style_config:
            for paragraph in cell.paragraphs:
                alignment_map = {
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                paragraph.alignment = alignment_map.get(style_config['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
        
        if 'vertical_alignment' in style_config:
            valign_map = {
                'center': WD_ALIGN_VERTICAL.CENTER,
                'top': WD_ALIGN_VERTICAL.TOP,
                'bottom': WD_ALIGN_VERTICAL.BOTTOM
            }
            cell.vertical_alignment = valign_map.get(style_config['vertical_alignment'], WD_ALIGN_VERTICAL.TOP)

    def _create_table(self, doc, table_config):
        """Crée un tableau avec configuration avancée"""
        headers = table_config.get('headers', [])
        rows = table_config.get('rows', [])
        merge_cells = table_config.get('merge_cells', [])
        styles = table_config.get('styles', {})
        
        # Créer le tableau
        num_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
        num_rows = len(rows) + (1 if headers else 0)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        
        # Style général du tableau
        table.style = styles.get('table_style', 'Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # En-têtes
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header)
                # Style header par défaut
                header_style = styles.get('header_style', {
                    'bold': True,
                    'background': 'D9E2F3',
                    'alignment': 'center',
                    'vertical_alignment': 'center'
                })
                self._apply_cell_style(header_cells[i], header_style)
        
        # Données
        start_row = 1 if headers else 0
        for i, row_data in enumerate(rows):
            row_cells = table.rows[start_row + i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = str(cell_data)
                
                # Style des cellules de données
                if 'cell_styles' in styles and f'{i},{j}' in styles['cell_styles']:
                    self._apply_cell_style(row_cells[j], styles['cell_styles'][f'{i},{j}'])
                elif 'data_style' in styles:
                    self._apply_cell_style(row_cells[j], styles['data_style'])
        
        # Fusion de cellules
        for merge in merge_cells:
            start_row = merge.get('start_row', 0)
            start_col = merge.get('start_col', 0)
            end_row = merge.get('end_row', start_row)
            end_col = merge.get('end_col', start_col)
            
            if start_row != end_row or start_col != end_col:
                start_cell = table.rows[start_row].cells[start_col]
                end_cell = table.rows[end_row].cells[end_col]
                start_cell.merge(end_cell)
        
        # Largeur des colonnes
        if 'column_widths' in styles:
            for i, width in enumerate(styles['column_widths']):
                for row in table.rows:
                    row.cells[i].width = Inches(width)
        
        return table

    def _parse_body_with_tables(self, doc, body_text):
        """Parse le corps du texte et insère les tableaux inline"""
        import re
        
        # Chercher les tableaux dans le texte
        table_pattern = r'\[TABLE\](.*?)\[/TABLE\]'
        parts = re.split(table_pattern, body_text, flags=re.DOTALL)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Texte normal
                self._add_text_content(doc, part)
            else:  # Tableau JSON
                try:
                    table_config = json.loads(part.strip())
                    self._create_table(doc, table_config)
                except json.JSONDecodeError as e:
                    doc.add_paragraph(f"[Erreur tableau: {str(e)}]")

    def _add_text_content(self, doc, text):
        """Ajoute du contenu texte (paragraphes et listes)"""
        if not text.strip():
            return
        
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Liste à puces
            if '- ' in para_text and para_text.split('\n')[0].strip().startswith('- '):
                lines = para_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('- '):
                        doc.add_paragraph(line[2:].strip(), style='List Bullet')
                    elif line:
                        doc.add_paragraph(line)
            else:
                doc.add_paragraph(para_text)

    def forward(self, title: str, recipient: str = "", sender: str = "", 
                date: str = "", subject: str = "", body: str = "", 
                tables: str = "", filename: str = "document", 
                style_config: str = ""):
        try:
            doc = Document()
            
            # === Configuration du style ===
            config = {}
            if style_config:
                try:
                    config = json.loads(style_config)
                except:
                    pass
            
            font_name = config.get('font', 'Arial')
            font_size = config.get('font_size', 11)
            margins = config.get('margins', 2.5)
            line_spacing = config.get('line_spacing', 1.15)
            
            # Marges
            section = doc.sections[0]
            section.top_margin = Cm(margins)
            section.bottom_margin = Cm(margins)
            section.left_margin = Cm(margins)
            section.right_margin = Cm(margins)
            
            # Style global
            style = doc.styles['Normal']
            style.font.name = font_name
            style.font.size = Pt(font_size)
            
            # Interligne
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = line_spacing
            
            # === En-tête lettre (optionnel) ===
            if sender:
                p = doc.add_paragraph(sender)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            if recipient:
                doc.add_paragraph(recipient)
            
            if date:
                doc.add_paragraph(date)
            
            if subject:
                p = doc.add_paragraph()
                p.add_run('Objet : ').bold = True
                p.add_run(subject)
                doc.add_paragraph()
            
            # === Titre ===
            if title:
                title_p = doc.add_paragraph(title)
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_p.runs[0]
                title_run.font.size = Pt(16)
                title_run.bold = True
                doc.add_paragraph()
            
            # === Corps avec tableaux inline ===
            if body:
                self._parse_body_with_tables(doc, body)
            
            # === Tableaux externes (en fin de document) ===
            if tables:
                try:
                    tables_array = json.loads(tables)
                    for table_config in tables_array:
                        doc.add_paragraph()
                        if 'title' in table_config:
                            p = doc.add_paragraph(table_config['title'])
                            p.runs[0].bold = True
                        self._create_table(doc, table_config)
                        doc.add_paragraph()
                except json.JSONDecodeError:
                    pass
            
            # === Formule de politesse (si format lettre) ===
            if recipient and sender:
                doc.add_paragraph()
                doc.add_paragraph("Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.")
                doc.add_paragraph()
                doc.add_paragraph(sender.split('\n')[0])  # Premier ligne = nom
            
            # === Sauvegarde ===
            safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).rstrip()
            file_path = f"{safe_filename}.docx"
            doc.save(file_path)
            
            return f"✓ Document Word créé avec succès : {os.path.abspath(file_path)}"
        
        except Exception as e:
            return f"✗ Erreur lors de la création du document : {str(e)}"

# class BuildWord(Tool):
#     name = "BuildWord"
#     description = (
#         "Crée un document Word professionnel (.docx) à partir d'un contenu structuré. "
#         "Utilise cette outil pour générer des lettres de motivation, CV, rapports, etc. "
#         "Le résultat sera un fichier Word avec mise en page propre (marges, police, interligne)."
#     )

#     inputs = {
#         "title": {"type": "string", "description": "Titre principal du document (ex: 'Lettre de Motivation')"},
#         "recipient": {"type": "string", "description": "Nom et adresse du destinataire (multiligne possible)"},
#         "sender": {"type": "string", "description": "Vos coordonnées (nom, adresse, email, tel)"},
#         "date": {"type": "string", "description": "Date (ex: 'Paris, le 23 décembre 2025')"},
#         "subject": {"type": "string", "description": "Objet de la lettre"},
#         "body": {"type": "string", "description": "Corps complet de la lettre en texte brut (paragraphes séparés par \\n\\n). Tu peux inclure des listes avec - item"},
#         "filename": {"type": "string", "description": "Nom du fichier sans extension (ex: 'Lettre_Motivation_Jonathan')"}
#     }

#     output_type = "string"

#     def forward(self, title: str, recipient: str, sender: str, date: str, subject: str, body: str, filename: str):
#         try:
#             doc = Document()

#             # === Configuration page professionnelle (format lettre française) ===
#             section = doc.sections[0]
#             section.top_margin = Cm(2.5)
#             section.bottom_margin = Cm(2.5)
#             section.left_margin = Cm(2.5)
#             section.right_margin = Cm(2.5)

#             # === Style global ===
#             style = doc.styles['Normal']
#             font = style.font
#             font.name = 'Arial'
#             font.size = Pt(11)

#             # === En-tête : Vos coordonnées (aligné à droite) ===
#             p = doc.add_paragraph(sender)
#             p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

#             # === Destinataire (aligné à gauche) ===
#             doc.add_paragraph(recipient)

#             # === Date (aligné à gauche ou droite selon norme) ===
#             doc.add_paragraph(date)

#             # === Objet ===
#             p = doc.add_paragraph()
#             p.add_run('Objet : ').bold = True
#             p.add_run(subject)

#             doc.add_paragraph()  # Espace

#             # === Titre centré ===
#             title_p = doc.add_paragraph(title)
#             title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             title_run = title_p.runs[0]
#             title_run.font.size = Pt(16)
#             title_run.bold = True

#             doc.add_paragraph()  # Espace

#             # === Corps de la lettre ===
#             paragraphs = body.split('\n\n')
#             for para_text in paragraphs:
#                 para_text = para_text.strip()
#                 if para_text.startswith('- '):
#                     # Liste à puces
#                     items = para_text.split('\n- ')
#                     for i, item in enumerate(items):
#                         if i == 0 and not item.startswith('- '):
#                             item = item[2:] if item.startswith('- ') else item
#                         p = doc.add_paragraph(item.strip('- '), style='List Bullet')
#                 else:
#                     doc.add_paragraph(para_text)

#             # === Formule de politesse ===
#             doc.add_paragraph()
#             doc.add_paragraph("Je vous prie d’agréer, Madame, Monsieur, l’expression de mes salutations distinguées.")

#             doc.add_paragraph()  # Espace signature
#             doc.add_paragraph("Jonathan Zadi")

#             # === Sauvegarde ===
#             safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).rstrip()
#             file_path = f"{safe_filename}.docx"
#             doc.save(file_path)

#             return f"Fichier Word professionnel créé avec succès : {os.path.abspath(file_path)} || {os.path.abspath(file_path)}"

#         except Exception as e:
#             return f"Erreur lors de la création du document Word : {str(e)}"


# class SendMail(Tool):
#     name = "send_mail"
#     description = (
#         "Envoie un e-mail avec support HTML et pièces jointes optionnelles. "
#         "Inputs: smtp_server, smtp_port, sender_email, sender_password, "
#         "recipient_email, subject, message, is_html (optionnel), attachment_path (optionnel)"
#     )

#     inputs = {
#         "smtp_server": {"type": "string", "description": "Adresse du serveur SMTP (ex: smtp.gmail.com)"},
#         "smtp_port": {"type": "number", "description": "Port SMTP (ex: 587 pour TLS, 465 pour SSL)"},
#         "sender_email": {"type": "string", "description": "Adresse email expéditeur (ex: davjonathan6@gmail.com)"},
#         "sender_password": {"type": "string", "description": "Mot de passe ou App Password de l'expéditeur (ex: qbcqkupoknwgeenf)"},
#         "recipient_email": {"type": "string", "description": "Adresse email destinataire (peut être une liste séparée par des virgules)"},
#         "subject": {"type": "string", "description": "Sujet du mail"},
#         "message": {"type": "string", "description": "Contenu du mail (texte ou HTML si is_html=True)"},
#         "is_html": {"type": "boolean", "description": "Si True, le message est interprété comme HTML (défaut: False)", "nullable": True},
#         "attachment_path": {"type": "string", "description": "Chemin vers un fichier à joindre (optionnel)", "nullable": True}
#     }

#     output_type = "string"

#     def forward(
#         self,
#         smtp_server: str,
#         smtp_port: int,
#         sender_email: str,
#         sender_password: str,
#         recipient_email: str,
#         subject: str,
#         message: str,
#         is_html: bool = False,
#         attachment_path: Optional[str] = None
#     ) -> str:
#         try:
#             # Validation
#             if not all([smtp_server, sender_email, sender_password, recipient_email, subject, message]):
#                 return "Erreur: Tous les champs obligatoires doivent être remplis."

#             # Création du mail
#             msg = MIMEMultipart()
#             msg["From"] = sender_email
#             msg["To"] = recipient_email
#             msg["Subject"] = subject

#             # Ajouter le message (texte ou HTML)
#             msg_type = "html" if is_html else "plain"
#             msg.attach(MIMEText(message, msg_type))

#             # Ajouter la pièce jointe si fournie
#             if attachment_path and os.path.isfile(attachment_path):
#                 with open(attachment_path, "rb") as attachment:
#                     part = MIMEBase('application', 'octet-stream')
#                     part.set_payload(attachment.read())
#                     encoders.encode_base64(part)
#                     part.add_header(
#                         'Content-Disposition',
#                         f'attachment; filename= {os.path.basename(attachment_path)}'
#                     )
#                     msg.attach(part)
#                 logger.info(f"Pièce jointe ajoutée: {attachment_path}")

#             # Connexion SMTP
#             if smtp_port == 465:
#                 # SSL
#                 server = smtplib.SMTP_SSL(smtp_server, smtp_port)
#             else:
#                 # TLS
#                 server = smtplib.SMTP(smtp_server, smtp_port)
#                 server.starttls()

#             server.login(sender_email, sender_password)

#             # Envoi
#             server.send_message(msg)
#             server.quit()

#             logger.info(f"Email envoyé de {sender_email} à {recipient_email}")
#             attachment_info = f" avec pièce jointe ({os.path.basename(attachment_path)})" if attachment_path else ""
#             return f"E-mail envoyé à {recipient_email} avec succès !{attachment_info}"

#         except FileNotFoundError:
#             return f"Erreur: Le fichier de pièce jointe '{attachment_path}' n'existe pas."
#         except smtplib.SMTPAuthenticationError:
#             return "Erreur: Échec de l'authentification. Vérifiez l'email et le mot de passe."
#         except smtplib.SMTPException as e:
#             return f"Erreur SMTP lors de l'envoi du mail : {e}"
#         except Exception as e:
#             logger.error(f"Erreur envoi email: {e}")
#             return f"Erreur lors de l'envoi du mail : {e}"

class BrevoSendMailDeprecated(Tool):
    name = "send_mail_brevo_deprecated"
    description = (
        "Envoie un email via Brevo (Sendinblue) avec support HTML "
        "et pièce jointe optionnelle."
    )

    inputs = {
        "recipient_email": {"type": "string", "description": "Email du destinataire"},
        "subject": {"type": "string", "description": "Sujet de l'email"},
        "message": {"type": "string", "description": "Contenu du message (HTML ou texte)"},
        "is_html": {"type": "boolean", "description": "Message HTML ?", "nullable": True},
        "attachment_path": {"type": "string", "description": "Chemin du fichier joint", "nullable": True}
    }

    output_type = "string"

    def forward(
        self,
        recipient_email: str,
        subject: str,
        message: str,
        is_html: bool = False,
        attachment_path: Optional[str] = None
    ) -> str:

        try:
            api_key = os.getenv("BREVO_API_KEY")
            sender_email = os.getenv("SENDER_EMAIL", "no-reply@agent-ia.com")

            if not api_key:
                return "Erreur : clé API Brevo manquante."

            configuration = sib_api_v3_sdk.Configuration()
            configuration.api_key['api-key'] = api_key

            api_instance = sib_api_v3_sdk.TransactionalEmailsApi(
                sib_api_v3_sdk.ApiClient(configuration)
            )

            email_data = {
                "to": [{"email": recipient_email}],
                "subject": subject,
                "sender": {"email": sender_email, "name": "candidAI"},
                "htmlContent": message if is_html else None,
                "textContent": message if not is_html else None,
            }

            # 📎 Pièce jointe
            if attachment_path and os.path.isfile(attachment_path):
                with open(attachment_path, "rb") as f:
                    encoded = base64.b64encode(f.read()).decode()

                email_data["attachment"] = [{
                    "content": encoded,
                    "name": os.path.basename(attachment_path)
                }]

            api_instance.send_transac_email(email_data)

            return f"📧 Email envoyé avec succès à {recipient_email}"

        except ApiException as e:
            return f"Erreur Brevo API : {e}"
        except Exception as e:
            return f"Erreur lors de l'envoi de l'email : {e}"


class SendMail(Tool):
    name = "send_mail"
    description = (
        "Envoie un email via SMTP Zoho avec support HTML "
        "et piece jointe optionnelle."
    )

    inputs = {
        "recipient_email": {"type": "string", "description": "Email du destinataire"},
        "subject": {"type": "string", "description": "Sujet de l'email"},
        "message": {"type": "string", "description": "Contenu du message HTML ou texte"},
        "is_html": {"type": "boolean", "description": "Message HTML ?", "nullable": True},
        "attachment_path": {"type": "string", "description": "Chemin du fichier joint", "nullable": True},
    }

    output_type = "string"

    def forward(
        self,
        recipient_email: str,
        subject: str,
        message: str,
        is_html: bool = False,
        attachment_path: Optional[str] = None,
    ) -> str:
        try:
            smtp_host = os.getenv("SMTP_HOST", "smtppro.zoho.com")
            smtp_port = int(os.getenv("SMTP_PORT", "465"))
            smtp_user = os.getenv("SMTP_USER") or os.getenv("SMTP_USERBOT")
            smtp_password = os.getenv("SMTP_PASSWORD")
            sender_email = os.getenv("SENDER_EMAIL") or smtp_user
            sender_name = os.getenv("SENDER_NAME", "IVOIR TRIPS INTERNATIONAL")

            if not all([smtp_host, smtp_port, smtp_user, smtp_password, sender_email]):
                return (
                    "Erreur SMTP: configuration email incomplete. "
                    "Verifiez SMTP_HOST, SMTP_PORT, SMTP_USERBOT, SMTP_PASSWORD et SENDER_EMAIL."
                )

            recipients = [
                email.strip()
                for email in recipient_email.split(",")
                if email.strip()
            ]
            if not recipients:
                return "Erreur SMTP: aucun destinataire valide."

            email_message = MIMEMultipart()
            email_message["From"] = formataddr((str(Header(sender_name, "utf-8")), sender_email))
            email_message["To"] = ", ".join(recipients)
            email_message["Subject"] = str(Header(subject, "utf-8"))
            email_message["Reply-To"] = sender_email

            body_type = "html" if is_html else "plain"
            email_message.attach(MIMEText(message, body_type, "utf-8"))

            if attachment_path and os.path.isfile(attachment_path):
                with open(attachment_path, "rb") as f:
                    attachment = MIMEBase("application", "octet-stream")
                    attachment.set_payload(f.read())
                encoders.encode_base64(attachment)
                attachment.add_header(
                    "Content-Disposition",
                    "attachment",
                    filename=os.path.basename(attachment_path),
                )
                email_message.attach(attachment)

            if smtp_port == 465:
                with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=30) as server:
                    server.login(smtp_user, smtp_password)
                    server.send_message(email_message)
            else:
                with smtplib.SMTP(smtp_host, smtp_port, timeout=30) as server:
                    server.ehlo()
                    server.starttls()
                    server.ehlo()
                    server.login(smtp_user, smtp_password)
                    server.send_message(email_message)

            return f"Email envoye avec succes a {', '.join(recipients)}"

        except smtplib.SMTPAuthenticationError:
            return (
                "Erreur SMTP: authentification refusee. "
                "Verifiez l'adresse Zoho et utilisez un mot de passe d'application si necessaire."
            )
        except smtplib.SMTPException as e:
            return f"Erreur SMTP lors de l'envoi de l'email : {e}"
        except Exception as e:
            logger.error(f"Erreur envoi email SMTP: {e}")
            return f"Erreur lors de l'envoi de l'email : {e}"
