from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from services.proforma_pdf import (
    ASSETS_DIR,
    CLIENT_DETAIL_FIELDS,
    DEFAULT_OUTPUT_DIR,
    TEAMBUILDING_AGENCY_FEE_RATE,
    _display_date,
    _format_amount_words_fcfa,
    _format_fcfa,
    _format_quantity,
    _format_rate_percent,
    _normalize_client_details,
    calculate_totals,
)


IVT_ORANGE_DARK = "EA580C"
IVT_ORANGE_SOFT = "FFF1E7"
IVT_ORANGE_FAINT = "FFF7ED"
IVT_INK = "101828"
IVT_MUTED = "667085"
PAGE_CONTENT_WIDTH = Mm(176)
WORD_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _add_watermark(header: Any) -> None:
    watermark_path = ASSETS_DIR / "watermark.png"
    if not watermark_path.exists():
        return
    try:
        relationship_id, _image = header.part.get_or_add_image(str(watermark_path))
        watermark = parse_xml(
            f"""
            <w:p
              xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
              xmlns:v="urn:schemas-microsoft-com:vml"
              xmlns:o="urn:schemas-microsoft-com:office:office"
              xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
              <w:r>
                <w:pict>
                  <v:shape
                    id="IvoirTripsWatermark"
                    o:spid="_x0000_s1025"
                    type="#_x0000_t75"
                    style="position:absolute;margin-left:167pt;margin-top:296pt;width:255pt;height:269pt;z-index:-251654144;mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
                    filled="f"
                    stroked="f">
                    <v:imagedata r:id="{relationship_id}" o:title="watermark"/>
                  </v:shape>
                </w:pict>
              </w:r>
            </w:p>
            """
        )
        header._element.append(watermark)
    except Exception:
        return


def _safe_word_path(reference: str, output_dir: str | Path | None = None) -> Path:
    clean_reference = "".join(
        char for char in str(reference or "proforma") if char.isalnum() or char in ("-", "_")
    ).strip("-_")
    if not clean_reference:
        clean_reference = "proforma"
    directory = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
    directory.mkdir(parents=True, exist_ok=True)
    return (directory / f"{clean_reference}.docx").resolve()


def get_proforma_word_path(reference: str, output_dir: str | Path | None = None) -> Path:
    return _safe_word_path(reference, output_dir)


def _set_font(
    run: Any,
    *,
    size: float,
    color: str = IVT_INK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Helvetica"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _format_paragraph(
    paragraph: Any,
    *,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 0,
) -> None:
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1


def _add_paragraph(
    document: Document,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    color: str = IVT_INK,
    size: float = 8.4,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    underline: bool = False,
    before: float = 0,
    after: float = 0,
) -> Any:
    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, alignment=alignment, before=before, after=after)
    run = paragraph.add_run(text)
    _set_font(run, size=size, color=color, bold=bold, italic=italic)
    run.underline = underline
    return paragraph


def _add_spacer(document: Document, height: float) -> None:
    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, after=height)


def _set_table_width(table: Any, width: Any = PAGE_CONTENT_WIDTH) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(width.twips))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_cell_width(cell: Any, width: Any) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(width.twips))


def _apply_column_widths(row: Any, widths: list[Any]) -> None:
    for cell, width in zip(row.cells, widths):
        _set_cell_width(cell, width)


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(
    cell: Any,
    *,
    top: int = 70,
    start: int = 90,
    bottom: int = 70,
    end: int = 90,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_borders(cell: Any, *, value: str = "nil", color: str = "FFFFFF", size: str = "0") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), value)
        tag.set(qn("w:color"), color)
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")


def _clear_table_borders(table: Any) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)


def _set_row_height(row: Any, height: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height.twips if hasattr(height, "twips") else height))
    tr_height.set(qn("w:hRule"), "atLeast")


def _cell_text(
    cell: Any,
    text: Any,
    *,
    bold: bool = False,
    italic: bool = False,
    color: str = IVT_INK,
    size: float = 8.2,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    fill: str | None = None,
    margins: tuple[int, int, int, int] = (70, 90, 70, 90),
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_margins(cell, top=margins[0], start=margins[1], bottom=margins[2], end=margins[3])
    if fill:
        _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    _format_paragraph(paragraph, alignment=alignment)
    lines = str(text or "").split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _set_font(run, size=size, color=color, bold=bold, italic=italic)


def _add_cell_paragraph(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    color: str = IVT_INK,
    size: float = 8.2,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    after: float = 0,
) -> Any:
    paragraph = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    _format_paragraph(paragraph, alignment=alignment, after=after)
    run = paragraph.add_run(text)
    _set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def _set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Mm(35)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(4)

    style = document.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(8.4)
    style.font.color.rgb = RGBColor.from_string(IVT_INK)

    header = section.header
    header.paragraphs[0].text = ""
    _format_paragraph(header.paragraphs[0])
    header.paragraphs[0].paragraph_format.left_indent = Mm(8)
    logo_path = ASSETS_DIR / "logo.png"
    if logo_path.exists():
        try:
            header.paragraphs[0].add_run().add_picture(str(logo_path), width=Mm(50), height=Mm(18))
        except Exception:
            pass
    _add_watermark(header)

    footer = section.footer
    footer.paragraphs[0].text = ""
    _format_paragraph(footer.paragraphs[0], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    line1 = footer.paragraphs[0].add_run(
        "Sarl au capital de 1.000.000 FCFA - Siège social : Cocody Angré 7è tranche - "
        "06 BP 914 ABIDJAN 06 RCC CI-ABJ-03-2021-"
    )
    _set_font(line1, size=6.8, color="000000")

    line2 = footer.add_paragraph()
    _format_paragraph(line2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run2 = line2.add_run("B13-02976 / NCC : 2148693 F")
    _set_font(run2, size=6.8, color="000000")

    line3 = footer.add_paragraph()
    _format_paragraph(line3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    email_run = line3.add_run("teambuilding@ivoirtrips.com / voyage@ivoirtrips.com")
    _set_font(email_run, size=6.8, color=IVT_ORANGE_DARK)
    phone_run = line3.add_run(" Tel : 07 79 181 778 / 05 95 298 183")
    _set_font(phone_run, size=6.8, color="000000")


def _add_document_header(document: Document, data: dict[str, Any], reference: str) -> None:
    title = _add_paragraph(
        document,
        "PROFORMA  N°",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=9,
        after=18,
    )
    ref_run = title.add_run(reference)
    _set_font(ref_run, size=9, bold=True)

    header_table = document.add_table(rows=1, cols=2)
    _set_table_width(header_table)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_column_widths(header_table.rows[0], [Mm(99), Mm(77)])
    _clear_table_borders(header_table)

    left_cell, right_cell = header_table.rows[0].cells
    left_cell.text = ""
    right_cell.text = ""
    _set_cell_margins(left_cell, top=0, start=0, bottom=0, end=0)
    _set_cell_margins(right_cell, top=0, start=0, bottom=0, end=0)
    _add_cell_paragraph(
        left_cell,
        "Centre des impôts : Cocody deux plateaux 3",
        italic=True,
        color=IVT_MUTED,
        size=7.4,
        after=14,
    )
    _add_cell_paragraph(
        left_cell,
        "Régime d'imposition : RSI",
        italic=True,
        color=IVT_MUTED,
        size=7.4,
    )

    _add_cell_paragraph(right_cell, "CLIENT", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_cell_paragraph(right_cell, str(data["client"]), bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    client_details = _normalize_client_details(data.get("client_details"))
    for detail_key, detail_label in CLIENT_DETAIL_FIELDS:
        detail_value = client_details.get(detail_key)
        if detail_value:
            _add_cell_paragraph(
                right_cell,
                f"{detail_label} : {detail_value}",
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )
    date_paragraph = _add_cell_paragraph(
        right_cell,
        "DATE : ",
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    date_run = date_paragraph.add_run(_display_date(data["date_proforma"]))
    _set_font(date_run, size=8.2, bold=True)

    _add_spacer(document, 14)
    _add_paragraph(
        document,
        str(data["objet"]).upper(),
        bold=True,
        underline=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        after=14,
    )


def _add_services_table(document: Document, sections: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=5)
    _set_table_width(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Mm(91), Mm(17), Mm(17), Mm(25), Mm(26)]

    headers = ["DESIGNATION", "JOURS", "QTE", "P.U. (FCFA)", "MONTANT HT"]
    _apply_column_widths(table.rows[0], widths)
    for cell, header in zip(table.rows[0].cells, headers):
        _cell_text(
            cell,
            header,
            bold=True,
            color="FFFFFF",
            size=7.4,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            fill=IVT_ORANGE_DARK,
        )
    _set_row_height(table.rows[0], Mm(8))

    for section in sections:
        section_row = table.add_row()
        _apply_column_widths(section_row, widths)
        merged = section_row.cells[0].merge(section_row.cells[-1])
        _cell_text(
            merged,
            str(section["nom"]).upper(),
            bold=True,
            color="FFFFFF",
            size=8.6,
            fill=IVT_ORANGE_DARK,
            margins=(70, 100, 70, 100),
        )
        _set_row_height(section_row, Mm(8))

        for item in section["prestations"]:
            row = table.add_row()
            _apply_column_widths(row, widths)
            values = [
                item["designation"],
                _format_quantity(item.get("nombre_jours")),
                _format_quantity(item.get("quantite")),
                _format_fcfa(item.get("prix_unitaire")) if item.get("prix_unitaire") else "",
                _format_fcfa(item.get("montant_ht")) if item.get("montant_ht") else "F CFA",
            ]
            for index, (cell, value) in enumerate(zip(row.cells, values)):
                _cell_text(
                    cell,
                    value,
                    size=8.2,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT,
                    margins=(75, 100, 75, 100),
                )
            _set_row_height(row, Mm(8))

        subtotal_row = table.add_row()
        _apply_column_widths(subtotal_row, widths)
        label_cell = subtotal_row.cells[0].merge(subtotal_row.cells[3])
        amount_cell = subtotal_row.cells[4]
        _cell_text(
            label_cell,
            f"TOTAL {str(section['nom']).upper()}",
            bold=True,
            color=IVT_ORANGE_DARK,
            size=8.3,
            fill=IVT_ORANGE_SOFT,
            margins=(95, 130, 95, 130),
        )
        _cell_text(
            amount_cell,
            _format_fcfa(section["sous_total"]),
            bold=True,
            size=8.4,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            fill=IVT_ORANGE_SOFT,
            margins=(95, 100, 95, 100),
        )
        _set_row_height(subtotal_row, Mm(9))

    _clear_table_borders(table)


def _add_financial_summary(
    document: Document,
    totals: dict[str, Any],
    agency_fee_label: str,
    vat_rate: Decimal,
) -> None:
    table = document.add_table(rows=5, cols=3)
    _set_table_width(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Mm(75), Mm(67), Mm(34)]
    for row in table.rows:
        _apply_column_widths(row, widths)

    rows = [
        ("RÉCAPITULATIF FINANCIER", "Sous-total mise en œuvre", totals["sous_total_ht"]),
        ("", agency_fee_label, totals["frais_agence"]),
        ("", "TOTAL HT", totals["total_ht"]),
        ("", f"TVA {format(vat_rate.normalize(), 'f')} % sur frais d’agence", totals["tva_frais_agence"]),
        ("", "TOTAL TTC", totals["total_ttc"]),
    ]
    for row_index, (title, label, amount) in enumerate(rows):
        row = table.rows[row_index]
        title_size = 14 if title else 8.2
        _cell_text(
            row.cells[0],
            title,
            bold=bool(title),
            color=IVT_ORANGE_DARK if title else IVT_INK,
            size=title_size,
            margins=(80, 0, 80, 0),
        )

        is_total_row = row_index == 4
        is_highlight_row = row_index in (2, 3)
        fill = IVT_ORANGE_DARK if is_total_row else IVT_ORANGE_SOFT if is_highlight_row else None
        text_color = "FFFFFF" if is_total_row else IVT_INK
        amount_color = "FFFFFF" if is_total_row else IVT_ORANGE_DARK if is_highlight_row else IVT_INK
        _cell_text(
            row.cells[1],
            label,
            bold=row_index >= 2,
            color=text_color,
            size=9.2,
            fill=fill,
            margins=(85, 100, 85, 100),
        )
        _cell_text(
            row.cells[2],
            _format_fcfa(amount),
            bold=row_index >= 2,
            color=amount_color,
            size=9.2 if row_index < 4 else 9.8,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            fill=fill,
            margins=(85, 100, 85, 100),
        )
        _set_row_height(row, Mm(8.5))

    _clear_table_borders(table)


def _add_amount_in_words(document: Document, amount: Any) -> None:
    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, after=8)
    label = paragraph.add_run("Montant arrêté à la somme de : ")
    _set_font(label, size=8.4, bold=True)
    value = paragraph.add_run(_format_amount_words_fcfa(amount))
    _set_font(value, size=8.4, color=IVT_ORANGE_DARK, bold=True)


def _add_info_box(document: Document, title: str, body: str) -> None:
    _add_spacer(document, 7)
    table = document.add_table(rows=1, cols=1)
    _set_table_width(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, IVT_ORANGE_FAINT)
    _set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _format_paragraph(paragraph)
    title_run = paragraph.add_run(title)
    _set_font(title_run, size=10, color=IVT_ORANGE_DARK, bold=True)
    paragraph.add_run().add_break()
    for line_index, line in enumerate(body.splitlines()):
        if line_index:
            paragraph.add_run().add_break()
        body_run = paragraph.add_run(line)
        _set_font(body_run, size=8.4)
    _clear_table_borders(table)


def _add_conditions(document: Document, totals: dict[str, Any], data: dict[str, Any]) -> None:
    document.add_page_break()
    _add_paragraph(
        document,
        "CONDITIONS COMMERCIALES & VALIDATION",
        bold=True,
        color=IVT_ORANGE_DARK,
        size=16,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
    )
    _add_amount_in_words(document, totals["total_ttc"])

    payment_terms = str(data.get("modalite_paiement") or "100 % à la commande").strip()
    payment_terms_text = payment_terms or "100 % à la commande"

    _add_info_box(
        document,
        "VALIDATION DE LA COMMANDE",
        "Merci de retourner la présente Proforma signée avec la mention "
        "« Bon pour Accord », ou de transmettre votre Bon de Commande (BC).",
    )
    _add_info_box(document, "MODALITÉS DE PAIEMENT", payment_terms_text.rstrip(".") + ".")
    _add_info_box(document, "PAIEMENT PAR CHÈQUE", "À l’ordre de : IVOIR TRIPS INTERNATIONAL")
    _add_info_box(
        document,
        "PAIEMENT PAR VIREMENT",
        "Bénéficiaire : IVOIR TRIPS INTERNATIONAL\n"
        "IBAN : CI93 CI008 01129 0129 467 319 90 22\n"
        "CODE SWIFT : SGCI CIAB",
    )

    _add_spacer(document, 16)
    signature_table = document.add_table(rows=1, cols=2)
    _set_table_width(signature_table)
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_column_widths(signature_table.rows[0], [Mm(88), Mm(88)])
    _cell_text(
        signature_table.rows[0].cells[0],
        "POUR LE CLIENT\nNom, signature et cachet\nMention « Bon pour Accord »",
        size=8.2,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        margins=(0, 0, 0, 0),
    )
    _cell_text(
        signature_table.rows[0].cells[1],
        "POUR IVOIR TRIPS INTERNATIONAL\nDirection Générale\nSignature et cachet",
        size=8.2,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        margins=(0, 0, 0, 0),
    )
    _clear_table_borders(signature_table)


def _add_control_line(
    document: Document,
    totals: dict[str, Any],
    agency_fee_label: str,
    vat_rate: Decimal,
) -> None:
    _add_spacer(document, 24)
    control_line = (
        f"Contrôle des calculs : sous-total mise en œuvre = {_format_fcfa(totals['sous_total_ht'])} ; "
        f"{agency_fee_label} = {_format_fcfa(totals['frais_agence'])} ; "
        f"total HT = {_format_fcfa(totals['total_ht'])} ; "
        f"TVA {format(vat_rate.normalize(), 'f')} % sur frais d’agence = "
        f"{_format_fcfa(totals['tva_frais_agence'])} ; "
        f"total TTC = {_format_fcfa(totals['total_ttc'])}."
    )
    _add_paragraph(document, control_line, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def generate_proforma_word(data: dict[str, Any], output_dir: str | Path | None = None) -> str:
    required_fields = ("reference", "client", "nombre_personnes", "date_proforma", "objet")
    missing = [field for field in required_fields if not data.get(field)]
    if missing:
        raise ValueError(f"Champs obligatoires manquants: {', '.join(missing)}")
    if int(data.get("nombre_personnes") or 0) <= 0:
        raise ValueError("nombre_personnes doit être supérieur à zéro.")

    agency_fee_rate = data.get("taux_frais_agence")
    if agency_fee_rate in (None, "") and str(data.get("pole") or "").lower() == "teambuilding":
        agency_fee_rate = TEAMBUILDING_AGENCY_FEE_RATE

    totals = calculate_totals(
        data.get("sections") or [],
        data.get("frais_agence") or 0,
        data.get("taux_tva_frais_agence") or 18,
        agency_fee_rate=agency_fee_rate,
    )
    sections = totals["sections"]
    if not sections:
        raise ValueError("La proforma doit contenir au moins une prestation.")

    reference = str(data["reference"]).strip()
    output_path = _safe_word_path(reference, output_dir)
    vat_rate = totals["taux_tva_frais_agence"]
    agency_rate_label = _format_rate_percent(totals.get("taux_frais_agence"))
    agency_fee_label = (
        f"Frais d’agence = {agency_rate_label} % du sous-total mise en œuvre"
        if agency_rate_label
        else "Frais d’agence"
    )

    document = Document()
    _set_document_defaults(document)
    _add_document_header(document, data, reference)
    _add_services_table(document, sections)
    _add_spacer(document, 12)
    _add_financial_summary(document, totals, agency_fee_label, vat_rate)
    _add_conditions(document, totals, data)
    _add_control_line(document, totals, agency_fee_label, vat_rate)
    document.save(output_path)
    return str(output_path)
